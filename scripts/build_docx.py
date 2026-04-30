#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_docx.py — Markdown 或 JSON spec → .docx 文档

依赖: python-docx (pip install python-docx)

使用:
    # Markdown 输入（最简单，推荐）
    python build_docx.py --md report.md --out report.docx
    python build_docx.py --md report.md --out report.docx --title "工艺方案"

    # JSON spec 输入（精细控制）
    python build_docx.py --spec doc.json --out report.docx

    # 从 stdin 读
    cat report.md | python build_docx.py --md - --out report.docx

    python build_docx.py --schema  # 打印 JSON spec 模板

JSON spec 格式:
{
  "title": "标题（可选，会变成大标题 + 居中）",
  "subtitle": "副标题",
  "author": "作者",
  "blocks": [
    {"type": "h1", "text": "..."},
    {"type": "h2", "text": "..."},
    {"type": "h3", "text": "..."},
    {"type": "p",  "text": "正文段落..."},
    {"type": "bullet", "items": ["...", "..."]},
    {"type": "ordered", "items": ["...", "..."]},
    {"type": "quote", "text": "..."},
    {"type": "code",  "text": "...", "lang": "python"},
    {"type": "table", "headers": ["A","B"], "rows": [["..","..."]]},
    {"type": "image", "path": "img.png", "caption": "..."},
    {"type": "pagebreak"}
  ]
}

Markdown 支持:
  # H1 / ## H2 / ### H3
  正文段落
  - 无序列表 / * 也行
  1. 有序列表
  > 引用
  ```code 块```
  | 表 | 头 |  ← GitHub 风格表格
  | -- | -- |
  ![alt](图片路径)
"""

import sys
import os
import io
import json
import argparse
import re
from pathlib import Path

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
else:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
except ImportError as e:
    print(f"[build_docx] ERROR: python-docx not installed.\n  pip install python-docx\n  Detail: {e}", file=sys.stderr)
    sys.exit(2)


CN_FONT = 'Microsoft YaHei'
EN_FONT = 'Calibri'


def set_run_font(run, size_pt=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = EN_FONT
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    # 设中文字体（python-docx 没有直接 API，用 oxml）
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), CN_FONT)
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)


def set_paragraph_default(paragraph, line_spacing=1.5, space_before_pt=0, space_after_pt=6):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)


# ---- Block builders ---------------------------------------------------------

def add_title_block(doc, title, subtitle=None, author=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_default(p, space_after_pt=4)
    r = p.add_run(title)
    set_run_font(r, size_pt=26, bold=True, color=(0x1A, 0x1F, 0x2C))

    if subtitle:
        ps = doc.add_paragraph()
        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_default(ps, space_after_pt=4)
        rs = ps.add_run(subtitle)
        set_run_font(rs, size_pt=14, color=(0x6B, 0x73, 0x80))

    if author:
        pa = doc.add_paragraph()
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_default(pa, space_after_pt=18)
        ra = pa.add_run(author)
        set_run_font(ra, size_pt=11, color=(0x6B, 0x73, 0x80))


def add_heading(doc, text, level=1):
    sizes = {1: 20, 2: 16, 3: 13}
    color = (0x1A, 0x1F, 0x2C) if level == 1 else (0x33, 0x33, 0x33)
    p = doc.add_paragraph()
    set_paragraph_default(p, space_before_pt=12, space_after_pt=4)
    r = p.add_run(text)
    set_run_font(r, size_pt=sizes.get(level, 12), bold=True, color=color)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_default(p)
    r = p.add_run(text)
    set_run_font(r, size_pt=11)


def add_list(doc, items, ordered=False):
    style = 'List Number' if ordered else 'List Bullet'
    for it in items:
        try:
            p = doc.add_paragraph(style=style)
        except KeyError:
            p = doc.add_paragraph()
            r = p.add_run(('1. ' if ordered else '• ') + str(it))
            set_run_font(r, size_pt=11)
            continue
        set_paragraph_default(p, space_after_pt=2)
        # python-docx adds bullets via the style; we just set font on the run
        for run in p.runs:
            set_run_font(run, size_pt=11)
        if not p.runs:
            r = p.add_run(str(it))
            set_run_font(r, size_pt=11)
        else:
            p.runs[0].text = str(it)


def add_quote(doc, text):
    p = doc.add_paragraph()
    set_paragraph_default(p, space_before_pt=6, space_after_pt=6)
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(text)
    set_run_font(r, size_pt=11, italic=True, color=(0x6B, 0x73, 0x80))


def add_code(doc, text):
    p = doc.add_paragraph()
    set_paragraph_default(p, space_before_pt=4, space_after_pt=4)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.4)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = 'Consolas'
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Consolas')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    r.font.color.rgb = RGBColor(0x2A, 0x2E, 0x38)


def add_table_block(doc, headers, rows):
    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    table = doc.add_table(rows=(1 if headers else 0) + len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    if headers:
        hdr = table.rows[0]
        for j, h in enumerate(headers):
            cell = hdr.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(h))
            set_run_font(r, size_pt=11, bold=True)
    for i, row in enumerate(rows):
        tr = table.rows[i + (1 if headers else 0)]
        for j, val in enumerate(row[:n_cols]):
            cell = tr.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, size_pt=10)
    # spacer
    doc.add_paragraph()


def add_image(doc, path, caption=None):
    if not os.path.exists(path):
        print(f"[build_docx] WARN: image not found: {path}", file=sys.stderr)
        add_paragraph(doc, f"[图片缺失: {path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    try:
        r.add_picture(path, width=Inches(5.5))
    except Exception as e:
        print(f"[build_docx] WARN: failed to add image {path}: {e}", file=sys.stderr)
        add_paragraph(doc, f"[图片加载失败: {path}]")
        return
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_default(cp, space_after_pt=10)
        cr = cp.add_run(caption)
        set_run_font(cr, size_pt=10, italic=True, color=(0x6B, 0x73, 0x80))


def add_pagebreak(doc):
    doc.add_page_break()


# ---- Markdown parser (lightweight, just enough for our needs) ---------------

def parse_markdown(md_text):
    blocks = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Heading
        m = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if m:
            blocks.append({'type': f'h{len(m.group(1))}', 'text': m.group(2).strip()})
            i += 1
            continue

        # Code fence
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            blocks.append({'type': 'code', 'text': '\n'.join(code_lines), 'lang': lang})
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            blocks.append({'type': 'quote', 'text': ' '.join(quote_lines)})
            continue

        # Unordered list
        if re.match(r'^\s*[-*+]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\s*[-*+]\s+', lines[i]):
                items.append(re.sub(r'^\s*[-*+]\s+', '', lines[i]))
                i += 1
            blocks.append({'type': 'bullet', 'items': items})
            continue

        # Ordered list
        if re.match(r'^\s*\d+\.\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                items.append(re.sub(r'^\s*\d+\.\s+', '', lines[i]))
                i += 1
            blocks.append({'type': 'ordered', 'items': items})
            continue

        # Image (single line)
        m = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', stripped)
        if m:
            blocks.append({'type': 'image', 'path': m.group(2), 'caption': m.group(1) or None})
            i += 1
            continue

        # Table
        if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers = [c.strip() for c in stripped.strip('|').split('|')]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
                i += 1
            blocks.append({'type': 'table', 'headers': headers, 'rows': rows})
            continue

        # Page break (--- or ***)
        if stripped in ('---', '***', '* * *'):
            blocks.append({'type': 'pagebreak'})
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Plain paragraph (consume until blank line)
        para = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,3}\s|>|```|[-*+]\s|\d+\.\s|\|)', lines[i].strip()):
            para.append(lines[i].strip())
            i += 1
        blocks.append({'type': 'p', 'text': ' '.join(para)})

    return blocks


# ---- Builder ----------------------------------------------------------------

def build_from_blocks(doc, blocks):
    for b in blocks:
        t = b.get('type')
        if t == 'h1':
            add_heading(doc, b.get('text', ''), level=1)
        elif t == 'h2':
            add_heading(doc, b.get('text', ''), level=2)
        elif t == 'h3':
            add_heading(doc, b.get('text', ''), level=3)
        elif t == 'p':
            add_paragraph(doc, b.get('text', ''))
        elif t == 'bullet':
            add_list(doc, b.get('items', []), ordered=False)
        elif t == 'ordered':
            add_list(doc, b.get('items', []), ordered=True)
        elif t == 'quote':
            add_quote(doc, b.get('text', ''))
        elif t == 'code':
            add_code(doc, b.get('text', ''))
        elif t == 'table':
            add_table_block(doc, b.get('headers', []), b.get('rows', []))
        elif t == 'image':
            add_image(doc, b.get('path', ''), b.get('caption'))
        elif t == 'pagebreak':
            add_pagebreak(doc)
        else:
            print(f"[build_docx] WARN: unknown block type '{t}'", file=sys.stderr)


def cli():
    ap = argparse.ArgumentParser(description="Markdown 或 JSON → .docx")
    ap.add_argument('--md', help="Markdown 文件路径（用 - 表示 stdin）")
    ap.add_argument('--spec', help="JSON spec 文件路径")
    ap.add_argument('--stdin-spec', action='store_true', help="从 stdin 读 JSON spec")
    ap.add_argument('--out', default='document.docx', help="输出 .docx")
    ap.add_argument('--title', help="文档标题（覆盖 spec 里的）")
    ap.add_argument('--subtitle')
    ap.add_argument('--author')
    ap.add_argument('--schema', action='store_true', help="打印 JSON spec 模板")
    args = ap.parse_args()

    if args.schema:
        sample = {
            "title": "工艺设计报告",
            "subtitle": "某某零件 2D 铣削方案",
            "author": "Manuscopy",
            "blocks": [
                {"type": "h1", "text": "1. 概述"},
                {"type": "p", "text": "本报告针对 6061-T6 铝合金法兰盘的批量化加工进行工艺设计。"},
                {"type": "h2", "text": "1.1 加工要求"},
                {"type": "bullet", "items": ["材料：6061-T6", "尺寸：Ø80 × 15 mm", "公差：±0.05"]},
                {"type": "h1", "text": "2. 切削参数"},
                {"type": "table",
                 "headers": ["工步", "刀具", "Vc(m/min)", "f(mm/r)", "ap(mm)"],
                 "rows": [["粗铣", "Φ12硬合金立铣刀", "180", "0.10", "1.5"],
                          ["精铣", "Φ8硬合金球头刀", "240", "0.05", "0.3"]]},
                {"type": "quote", "text": "粗加工求效率，精加工求精度。"},
                {"type": "pagebreak"},
                {"type": "h1", "text": "3. NC 程序"},
                {"type": "code", "lang": "gcode", "text": "G90 G54 G17\nG00 X0 Y0 Z5\n..."},
            ]
        }
        print(json.dumps(sample, ensure_ascii=False, indent=2))
        return

    title = args.title
    subtitle = args.subtitle
    author = args.author
    blocks = []

    if args.md:
        if args.md == '-':
            md_text = sys.stdin.read()
        else:
            with open(args.md, 'r', encoding='utf-8') as f:
                md_text = f.read()
        blocks = parse_markdown(md_text)
        # Auto-extract title from first h1 if not given
        if not title:
            for b in blocks:
                if b.get('type') == 'h1':
                    title = b['text']
                    blocks = [x for x in blocks if x is not b]  # remove that h1
                    break
    elif args.spec or args.stdin_spec:
        if args.stdin_spec:
            spec = json.load(sys.stdin)
        else:
            with open(args.spec, 'r', encoding='utf-8') as f:
                spec = json.load(f)
        title = title or spec.get('title')
        subtitle = subtitle or spec.get('subtitle')
        author = author or spec.get('author')
        blocks = spec.get('blocks', [])
    else:
        ap.error("provide --md FILE or --spec FILE (or --stdin-spec). Use --schema to see the format.")

    doc = Document()
    # Default style
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = EN_FONT

    if title:
        add_title_block(doc, title, subtitle, author)

    build_from_blocks(doc, blocks)

    out_path = Path(args.out).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    size = out_path.stat().st_size
    print(f"OK: {out_path} ({size/1024:.1f} KB, {len(blocks)} blocks)")


if __name__ == '__main__':
    cli()
